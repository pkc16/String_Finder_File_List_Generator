"""
Application to:
  - search a given directory for occurrences of search string in files
  - generate list of all files in a given directory 

File types searched: .txt, .docx, .pdf
does not search .doc files because those old Word files cause errors
PDF search will search actual text, not text which are in image files in the PDF
if exceptions encountered during file operations, file is skipped and reported as being skipped

if memory problems, try using mmap:
      https://stackoverflow.com/questions/4940032/how-to-search-for-a-string-in-text-files

future enhancements:
    - add minimal characters restriction in search string
    - test Word files (test file which has regular text and table; test paragraphs with different styles)
"""

from tkinter import *
from tkinter import filedialog, messagebox
import os
import re
import pathlib
import docx
import PyPDF2

class StringFinder(object):
    def __init__(self, window):
        self.window = window
        self.window.wm_title("String Finder / File List Generator Application")
        self.directory = ""
        self.outputFileName = "0_String_Finder_results.txt"
        self.outputFilePath = ""
        self.varChkTxt = IntVar()
        self.varChkDocx = IntVar()
        self.varChkPDF = IntVar()
        self.createFields()


    def createFields(self):
        self.lblEmpty = Label(self.window, text="")
        self.lblEmpty.grid(row=0, column=0)

        self.lblDir = Label(self.window, text=self.directory)
        self.lblDir.grid(row=1, column=2, columnspan=2, sticky=W)

        self.lblSearch = Label(self.window, text="Enter search string: ", padx=5, pady=5)
        self.lblSearch.grid(row=2, column=1)

        self.searchStrEntry = Entry(self.window, width=50)
        self.searchStrEntry.grid(row=2, column=2)

        self.btnSearch = Button(self.window, text="Search", command=lambda: self.search(self.searchStrEntry.get()), state=DISABLED)
        self.btnSearch.grid(row=2, column=3, padx=10, pady=10)

        self.btnExit = Button(self.window, text="Exit", command=self.window.destroy)
        self.btnExit.grid(row=2, column=4)

        self.btnBrowse = Button(self.window, text="Browse", command=self.getDirectory)
        self.btnBrowse.grid(row=1, column=1)

        self.btnFileStructure = Button(self.window, text="Output File List", command=self.outputDirectoryStructure)
        self.btnFileStructure.grid(row=1, column=3)

        self.lblResult = Label(self.window, text="")
        self.lblResult.grid(row=7, column=1, columnspan=2) 

        self.lblBypassedFiles = Label(self.window, text="")
        self.lblBypassedFiles.grid(row=8, column=1)

        self.lblOutputFile = Label(self.window, text="")
        self.lblOutputFile.grid(row=9, column=1)

        self.lblIncludeFileTypes = Label(self.window, text="Search on:")
        self.lblIncludeFileTypes.grid(row=3, column=1, sticky=W)

        self.chkTxt = Checkbutton(self.window, text=".txt", variable=self.varChkTxt)
        self.chkTxt.grid(row=4, column=1, sticky=W)

        self.chkDocx = Checkbutton(self.window, text=".docx", variable=self.varChkDocx)
        self.chkDocx.grid(row=5, column=1, sticky=W)

        self.chkPDF = Checkbutton(self.window, text=".pdf", variable=self.varChkPDF)
        self.chkPDF.grid(row=6, column=1, sticky=W)


    def getDirectory(self):
        # get directory to search
        os.chdir("C:\\")
        path = filedialog.askdirectory(initialdir="C:", title="Select a directory to search")
        os.chdir(path)

        self.directory = path
        self.lblDir.grid_forget()
        self.lblDir = Label(self.window, text=self.directory)
        self.lblDir.grid(row=1, column=2, columnspan=2, sticky=W)
        self.searchStrEntry.grid_forget()
        self.searchStrEntry = Entry(self.window, width=50)
        self.searchStrEntry.grid(row=2, column=2)
        self.btnSearch = Button(self.window, text="Search", command=lambda: self.search(self.searchStrEntry.get()))
        self.btnSearch.grid(row=2, column=3, padx=10, pady=10)
        self.outputFilePath = os.path.join(self.directory, self.outputFileName)
        self.lblResult.grid_forget()
        self.lblOutputFile.grid_forget()
        

    def getListOfFiles(self, dirName):
        # Create a list of all files in the passed directory including subdirectories
        listOfFile = os.listdir(dirName)
        allFiles = []
        # Iterate over all the entries
        for entry in listOfFile:
            # Create full path
            fullPath = os.path.join(dirName, entry)
            # If entry is a directory then get the list of files in this directory 
            if os.path.isdir(fullPath):
                allFiles = allFiles + self.getListOfFiles(fullPath)
            else:
                allFiles.append(fullPath)
                    
        return allFiles      


    def outputDirectoryStructure(self):
        # output list of all files inside directory
        # part 1 lists files with paths; part 2 lists filenames only
        listAllFiles = []
        listAllFiles = self.getListOfFiles(self.directory)

        listAllFiles = [file.replace('\\', '/') for file in listAllFiles]
        
        # Remove any files from the list which have no suffix or are hidden files
        listAllFiles = [file for file in listAllFiles if pathlib.Path(file).suffix[1:] != '' and not file.startswith('.')]
        
        outfilepath = os.path.join(self.directory, "0_File_list.txt")

        with open(outfilepath, "w") as outFile:
            outFile.write("1. Files in " + self.directory + ":\n\n")
            for file in listAllFiles:
                outFile.write(file + "\n")

            # now write the filenames without the path
            outFile.write("\n\n")
            outFile.write("2. Filenames only:\n\n")
            for file in listAllFiles:
                outFile.write(os.path.basename(file) + "\n")


    def search(self, searchStr):
        # Search directory for the string
        if self.directory == "":
            messagebox.showwarning("Directory Not Specified", "Please specify the directory to search in.")
            return
        elif self.searchStrEntry.get() == "":
            messagebox.showwarning("Search String Not Specified", "Please enter the search string.")
            return
        elif self.varChkTxt.get() == 0 and self.varChkDocx.get() == 0 and self.varChkPDF.get() == 0:
            messagebox.showwarning("File Types Not Specified", "Please select at least 1 file type to search on.")
            return

        # Get list of all files in the directory
        listAllFiles = []
        listAllFiles = self.getListOfFiles(self.directory)
    
        # Need to replace "\\" with "/" since above function with prefix filename with "\\"
        listAllFiles = [file.replace('\\', '/') for file in listAllFiles]

        # Get the file types to search on
        typeList = []
        if self.varChkTxt.get() == 1:
            typeList.append('txt')

        if self.varChkDocx.get() == 1:
            typeList.append('docx')

        if self.varChkPDF.get() == 1:
            typeList.append('pdf')

        # Remove any files from the list which are not in the file types or have no suffix or are hidden files
        listAllFiles = [file for file in listAllFiles if pathlib.Path(file).suffix[1:] in typeList and pathlib.Path(file).suffix[1:] != '' and not file.startswith('.')]

        # Now loop through files and find those which match the search string
        foundInFilesCount = 0
        numBypassedFiles = 0
        filesBypassed = []
        with open(self.outputFilePath, "w") as outFile:
            
            outFile.write("\"" + searchStr + "\" found in the following files in " + self.directory + ":\n\n")
            for txtfile in listAllFiles:
                suffix = pathlib.Path(txtfile).suffix[1:]

                if suffix == 'txt':
                    try:
                        with open(txtfile) as f:
                            if searchStr.lower() in f.read().lower():
                                outFile.write(txtfile + "\n")
                                foundInFilesCount += 1

                    except:
                        filesBypassed.append(txtfile)
                        numBypassedFiles += 1
                        continue

                elif suffix == 'docx':
                    try:
                        doc = docx.Document(txtfile)

                        #regex = re.compile(searchStr.lower())
                        for p in doc.paragraphs:
                            if p.text.lower().find(searchStr.lower()) >= 0:
                                outFile.write(txtfile + "\n")
                                foundInFilesCount += 1
                                break
                        # else:
                        #     continue
                        # break

                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.lower().find(searchStr.lower()) >= 0:
                                        outFile.write(txtfile + "\n")
                                        foundInFilesCount += 1
                                        break
                                else:
                                    continue
                                break
                            else:
                                continue
                            break

                    except:
                        filesBypassed.append(txtfile)
                        numBypassedFiles += 1
                        continue
                
                elif suffix == 'pdf':
                    try:
                        pdfFile = open(txtfile, 'rb')
                        reader = PyPDF2.PdfFileReader(pdfFile)
                        for pageNum in range(reader.numPages):
                            text = reader.getPage(pageNum).extractText()
                            if searchStr.lower() in text.lower():
                                outFile.write(txtfile + "\n")
                                foundInFilesCount += 1
                                break
                        pdfFile.close()

                    except:
                        # likely exception is PyPDF2.utils.PdfReadError; occurs if pdf is encrypted (PyPDF2 can't read encrypted file)
                        pdfFile.close()
                        filesBypassed.append(txtfile)
                        numBypassedFiles += 1
                        continue

            if numBypassedFiles > 0:
                outFile.write("\n\nFiles skipped due to error trying to read file:\n")
                outFile.write("(check these files individually)\n\n")
                for item in filesBypassed:
                    outFile.write(item + "\n")

        self.lblResult.grid_forget()
        if foundInFilesCount == 0 and numBypassedFiles == 0:
            self.lblResult = Label(self.window, text="Search string was not found")
            self.lblResult.grid(row=7, column=1, columnspan=2, sticky=W)
            self.lblOutputFile.grid_forget()

            # delete the output file since search did not find any matches
            os.unlink(self.outputFilePath)
            return
        else:
            self.lblResult = Label(self.window, text=f"Number of files found in: {str(foundInFilesCount)}")
            self.lblResult.grid(row=7, column=1, columnspan=2, sticky=W)
        
        self.lblBypassedFiles.grid_forget()
        self.lblBypassedFiles = Label(self.window, text=f"Number of files bypassed: {str(numBypassedFiles)}")
        self.lblBypassedFiles.grid(row=8, column=1, columnspan=2, sticky=W)

        self.lblOutputFile.grid_forget()
        self.lblOutputFile = Label(self.window, text="Results file: " + self.outputFileName)
        self.lblOutputFile.grid(row=9, column=1, columnspan=2, sticky=W)



root = Tk()
simple_app = StringFinder(root)

root.geometry("600x300+300+350")  #(window width x window height + position right + position down)

root.mainloop()